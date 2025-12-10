"""
Text extraction from various file formats (PDF, DOCX, images)
"""
import io
import base64
from typing import Optional
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
import os

class TextExtractor:
    """Extract text from different file formats"""
    
    @staticmethod
    def extract_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            extracted_text = "\n".join(text_parts)
            
            # If no text extracted, might be scanned PDF - use OCR
            if not extracted_text.strip():
                return TextExtractor.extract_from_pdf_with_ocr(file_content)
            
            return extracted_text.strip()
            
        except Exception as e:
            print(f"PDF extraction error: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_pdf_with_ocr(file_content: bytes) -> str:
        """
        Extract text from scanned PDF using OCR
        """
        try:
            from pdf2image import convert_from_bytes
            
            # Convert PDF to images
            images = convert_from_bytes(file_content)
            
            text_parts = []
            for image in images:
                # Use pytesseract for OCR
                text = pytesseract.image_to_string(image)
                if text:
                    text_parts.append(text)
            
            return "\n".join(text_parts).strip()
            
        except Exception as e:
            print(f"PDF OCR extraction error: {e}")
            raise Exception(f"Failed to extract text from scanned PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts).strip()
            
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_from_image(file_content: bytes) -> str:
        """
        Extract text from image using OCR
        """
        try:
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text.strip()
            
        except Exception as e:
            print(f"Image OCR extraction error: {e}")
            raise Exception(f"Failed to extract text from image: {str(e)}")
    
    @staticmethod
    def extract_from_base64(base64_content: str, file_type: str) -> str:
        """
        Extract text from base64 encoded file
        """
        try:
            # Remove data URL prefix if present
            if ',' in base64_content:
                base64_content = base64_content.split(',')[1]
            
            # Decode base64
            file_content = base64.b64decode(base64_content)
            
            # Route to appropriate extractor
            file_type = file_type.lower()
            
            if file_type in ['pdf', 'application/pdf']:
                return TextExtractor.extract_from_pdf(file_content)
            elif file_type in ['docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return TextExtractor.extract_from_docx(file_content)
            elif file_type in ['png', 'jpg', 'jpeg', 'image/png', 'image/jpeg', 'image/jpg']:
                return TextExtractor.extract_from_image(file_content)
            else:
                raise Exception(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            print(f"Base64 extraction error: {e}")
            raise Exception(f"Failed to extract text from file: {str(e)}")
    
    @staticmethod
    def extract_text(file_content: bytes, filename: str) -> str:
        """
        Auto-detect file type and extract text
        """
        try:
            # Determine file type from extension
            file_ext = filename.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                return TextExtractor.extract_from_pdf(file_content)
            elif file_ext in ['docx', 'doc']:
                return TextExtractor.extract_from_docx(file_content)
            elif file_ext in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                return TextExtractor.extract_from_image(file_content)
            else:
                raise Exception(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            print(f"Text extraction error: {e}")
            raise Exception(f"Failed to extract text: {str(e)}")
